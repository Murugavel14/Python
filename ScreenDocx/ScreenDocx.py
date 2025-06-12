import customtkinter as ctk
import pyautogui
from PIL import Image, ImageDraw, ImageFont
from tkinter import filedialog, Toplevel, Label
from docx import Document
from docx.shared import Inches
import datetime, os, tempfile, re
import pygame

# Appearance Section
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("dark-blue")

# Globals
screenshot_files = []
save_folder = None
camera_sound = None

#sound init
pygame.mixer.init()
if os.path.exists("mouse.mp3"):
    try:
        camera_sound = pygame.mixer.Sound("mouse.mp3") 
        camera_sound.set_volume(1.0)  # Adjust volume here
    except pygame.error:
        camera_sound = None

# Font Section
font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" # if use linux this directory for font
if not os.path.exists(font_path):
    font_path = "arial.ttf"  #it's windows users

# Tooltip Section
class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip_window = None
        self.after_id = None
        widget.bind("<Enter>", self.schedule_show)
        widget.bind("<Leave>", self.cancel_show)

    def schedule_show(self, event=None):
        self.after_id = self.widget.after(500, self.show)

    def cancel_show(self, event=None):
        if self.after_id:
            self.widget.after_cancel(self.after_id)
            self.after_id = None
        self.hide()

    def show(self):
        if self.tip_window:
            return
        x = self.widget.winfo_rootx() + 40
        y = self.widget.winfo_rooty() + 10
        self.tip_window = tw = Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.configure(cursor="arrow")
        tw.geometry(f"+{x}+{y}")
        label = Label(tw, text=self.text, bg="#333", fg="white", font=("Arial", 10), relief="solid", bd=1)
        label.pack()

    def hide(self):
        if self.tip_window:
            self.tip_window.destroy()
            self.tip_window = None

# Folder chooser Section
def choose_folder():
    global save_folder, screenshot_files
    folder = filedialog.askdirectory()
    if folder:
        save_folder = folder
        screenshot_files.clear()
        btn_ss.configure(state="normal")
        btn_save.configure(state="normal")

# Screenshot logic Section
def take_screenshot():
    now = datetime.datetime.now()
    filename = now.strftime("screenshot_%H%M%S.jpg")
    path = os.path.join(tempfile.gettempdir(), filename)

    app.withdraw()  # Hide window fully
    app.after(200, lambda: complete_screenshot(now, path)) #if you taking Screenshot to fliker the ScreenDocx window

def complete_screenshot(now, path):
    img = pyautogui.screenshot()

    draw = ImageDraw.Draw(img)
    timestamp = now.strftime("%Y-%m-%d %H:%M:%S")
    font = ImageFont.truetype(font_path, 20)
    draw.text((10, 10), timestamp, fill="white", font=font)

    img = img.resize((img.width // 2, img.height // 2))
    img = img.convert("RGB")
    img.save(path, "JPEG", quality=100)

    screenshot_files.append(path)

    if sound_toggle.get() and camera_sound:
        camera_sound.play()

    app.deiconify()  # Show window again instantly

# Save and create auto-numbered DOCX
def save_and_exit():
    global save_folder
    if not save_folder or not screenshot_files:
        return

    existing_files = [f for f in os.listdir(save_folder) if re.match(r'ScreenDocx(\d+)\.docx$', f, re.IGNORECASE)]
    numbers = [int(re.search(r'(\d+)', name).group(1)) for name in existing_files]
    next_num = max(numbers, default=0) + 1
    doc_name = f"ScreenDocx{next_num}.docx"
    doc_path = os.path.join(save_folder, doc_name)

    doc = Document()
    for shot in screenshot_files:
        doc.add_picture(shot, width=Inches(6))
        doc.add_paragraph("")

    doc.save(doc_path)
    print(f"✅ Saved to {doc_name}")

    # ✅ Delete images after saving
    for shot in screenshot_files:
        if os.path.exists(shot):
            os.remove(shot)

    screenshot_files.clear()  # ✅ Important to avoid duplication!
    btn_ss.configure(state="disabled")
    btn_save.configure(state="disabled")
    save_folder = None

# GUI setup
app = ctk.CTk()
app.geometry("325x100")
app.title("ScreenDocx")
app.resizable(False, False)
#app.iconbitmap("Screendocx.ico")

frame = ctk.CTkFrame(app, fg_color="#1f1f1f", corner_radius=15)
frame.pack(pady=20, padx=10, fill="both", expand=True)

btn_folder = ctk.CTkButton(frame, text="📁", width=40, height=40, corner_radius=30, command=choose_folder, font=("Arial", 22))
btn_folder.grid(row=0, column=0, padx=10) # Button for File choose

btn_ss = ctk.CTkButton(frame, text="📸", width=40, height=40, corner_radius=30, command=take_screenshot, font=("Arial", 22), state="disabled")
btn_ss.grid(row=0, column=1, padx=10) # Button for Taking Screenshot

btn_save = ctk.CTkButton(frame, text="📝", width=40, height=40, corner_radius=30, command=save_and_exit, font=("Arial", 22), state="disabled")
btn_save.grid(row=0, column=2, padx=10) # Button for Saving Screenshot 

sound_toggle = ctk.CTkCheckBox(frame, text="🔊", font=("Arial", 12), checkbox_height=18, checkbox_width=18)
sound_toggle.grid(row=0, column=3, padx=10) # Button for Sound
sound_toggle.select()

# Tooltips
ToolTip(btn_folder, "Choose folder")
ToolTip(btn_ss, "Take Screenshot")
ToolTip(btn_save, "Save to DOCX")

app.mainloop()

#Code BY Murugavel
